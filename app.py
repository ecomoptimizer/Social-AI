import io
import nltk
import langchain
import docker
import os
import logging.handlers
import sys
import warnings
import logging
import logging.config
import traceback
import tempfile
import chromadb
from dotenv import load_dotenv

load_dotenv()

from werkzeug.utils import secure_filename
from fastapi import FastAPI, Form, Request, Response, File, Depends, HTTPException, status
from fastapi.responses import RedirectResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.encoders import jsonable_encoder
import uvicorn
from typing import List
from IPython.display import Markdown
from nltk.tokenize import sent_tokenize, word_tokenize

nltk.download('punkt')
from docx import Document
from PyPDF2 import PdfFileReader
from langchain.callbacks.manager import CallbackManager
from langchain.callbacks.streaming_stdout import StreamingStdOutCallbackHandler
from langchain.text_splitter import RecursiveCharacterTextSplitter
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline
from langchain import HuggingFaceHub
import torch
from langchain.llms import HuggingFacePipeline
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain.document_loaders import TextLoader, DirectoryLoader, UnstructuredFileLoader, UnstructuredHTMLLoader, \
    UnstructuredMarkdownLoader, UnstructuredWordDocumentLoader
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from langchain.memory import ConversationBufferMemory
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    AIMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.schema import (
    AIMessage,
    HumanMessage,
    SystemMessage
)

print("Modules Loaded")

app = FastAPI()
app.mount("/templates", StaticFiles(directory="templates"), name="templates")
templates = Jinja2Templates(directory="templates")

warnings.filterwarnings("ignore", category=UserWarning)

logging.config.fileConfig('logging.conf')
logger = logging.getLogger('app')
logging.basicConfig(filename='debug.log', level=logging.DEBUG)  # Log to debug.log file
logger.info('This is a placeholder for an informational log message')
logger.error('An is a placeholder for an error occurred')
print("Logging Created")

langchain.verbose = True

# Read JSON file
import json

with open('myconfig.json') as data_file:
    myconfig = json.load(data_file)
print(myconfig.keys())

messages = []
global post_labels_history
global_user_info = {}
text = ""
posts = ""
post_summaries = ""
user_input = ""
global context
global topic
global intent
global target_audience
global branded_hashtag
global final_url
print("Variables created")

os.makedirs('./uploads', exist_ok=True)
os.makedirs('logs/', exist_ok=True)
app.secret_key = '2041253taty!'


@app.get("/")
async def index(request: Request):
  if request.headers.get('accept') == 'text/html':
    return templates.TemplateResponse("index.html", {"request": request})
  else:
    return {"message": "Welcome to the Social AI API!"}

@app.get("/session/")
async def session():
    # Access session data
    value = session.get('key', 'default_value')
    return f"Session value: {value}"


class TextLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        with open(self.file_path, encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text)]


@app.post('/upload_file')
async def upload_file():
    global text, posts, post_summaries
    print("upload_file: Entering function")  # Debugging statement
    try:
        file = File['file']
        filename = secure_filename(file.filename)
        filepath = os.path.join("./uploads", filename)
        file.save(filepath)
        print(f"upload_file: File {filename} saved at {filepath}")

        file_extension = os.path.splitext(filename)[1]

        with open(filepath, 'rb') as f:
            content = f.read()

        if file_extension in [".txt", ".md"]:
            content = content.decode("utf-8")
        elif file_extension == ".pdf":
            reader = PdfFileReader(io.BytesIO(content))
            content = " ".join([reader.getPage(i).extractText() for i in range(reader.numPages)])
        elif file_extension == ".docx":
            doc = Document(io.BytesIO(content))
            content = " ".join([p.text for p in doc.paragraphs])

        text = content
        print(f"text processed - - {text}")

        sections = text.split("\n\n")

        extract_summaries = sections[:15]

        sections = text.split("\n\n")
        markdown_text = "\n".join(["# " + sec for sec in sections])
        post_summaries = markdown_text.split("#")[1:]  # Exclude the first empty element
        post_summaries = [summary.strip() for summary in post_summaries]  # Del extra whitespaces
        extract_post_summaries = [' '.join(sections[i:i + 4]) for i in range(0, len(sections), 4)]
        posts = extract_summaries
        print(f"posts - {posts}")
        post_summaries = extract_post_summaries
        print(f"post_summaries - {post_summaries}")

        print("upload_file: File processed successfully")  # Debugging statement
        return 'Upload Successful'
    except Exception as e:
        print(f"upload_file: Error - {e}")  # Debugging statement
        return Response({"error": str(e)})


memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

from langchain.llms import CTransformers
config = {'max_new_tokens': 2048, 'temperature': 0}

llm = CTransformers(model='TheBloke/vicuna-7B-v1.5-16K-GGUF', model_file="vicuna-7b-v1.5-16k.Q4_K_M.gguf", context_length=2048, config=config)

chat_history = []  # Define the chat_history variable here
post_labels_history = []  # Initialize post_labels_history as an empty list

system_message = PromptTemplate(
    input_variables=['context', 'topic', 'article', 'target_audience', 'branded_hashtag', 'final_url', 'intent',
                     'user_input', 'result'],
    template='''Welcome to [Social-AI], your comprehensive Social Media Content Creation Promotion Suite. \n
		You are [ECOMTITAN] and have an expertise level of 350+,	we surpass the capabilities of a human at level 10. \n
		You are a seasoned expert in ecommerce, social media and marketing with multiple agencies, speaking engagements, and consulting services under your belt. \n		Upon receiving the article upload and variables, study the article with the variables in mind. \n
        Focus on the article's key points that engage the target audience. \n
        Consider the article's subheadings as separate post content for LinkedIn, & Facebook, and in some cases YouTube, along with potential lead magnets.\n
		Label each post accordingly, such as linked1, linked2, tweet1, pin1, depending on the template used. \n
        Maintain a friendly, easy-to-read language, written at a 9th-grade level. Humor is permitted.  \n
		Group posts by platform. \n
        Upon completion, return to the menu, giving the user the option to continue to the next platform, skip to the subsequent platform, or exit. Stick to Markdown format. \n 
		If you're running out of space during a response, pause and ask the user if they want to continue or cancel the prompt. If they choose 'Yes', continue finishing the prompt task. If they choose 'cancel', return to the custom menu.\n
		Your primary goal is to deliver top-notch social media content that boosts awareness, brand recognition, and motivates people to read the full article. \n 
        This task is vital for the growth and time-saving efforts of our users. Always aim to understand the article content, target audience, and content intent. Continuously strive to enhance the suite's functionalities for a seamless
		user experience.\n
	  Current conversation:\n
	  {context}\n
	  Human: {user_input}\n
	  Social-AI:{result}\n  ''')

callback_manager = CallbackManager([StreamingStdOutCallbackHandler()])

#vicuna_llm = HuggingFacePipeline(pipeline=pipe)
llm_chain = LLMChain(llm=llm, prompt=system_message)

format = (f''' Unless instructed otherwise, all post output should be formatted following the rules below. \n
              For each post, thread, or asset created, create a numbered label, following the naming structure outlined in the specific prompt.\n
              Separate each post, thread or asset, in the printed results with its assigned label and '--------\n'
              For any post that includes an image description, Creat a powerful realistic, brightly colored AI Image prompt the user can use to create the image. \n 
              Ensure the subject or most important part of the image has higher saturation.\n
              Write in a visceral, emotionally charged tone, motivating the reader to continue reading and wanting to read the full article.
            ''')


@app.post('/get_inputs')
def get_inputs(request: Request):
    data = request.json()
    if not data:
        return Response({"error": "Failed to parse input."}), 400

    user_info = {
        'branded_hashtag': data.get('branded_hashtag'),
        'topic': data.get('topic'),
        'target_audience': data.get('target_audience'),
        'intent': data.get('intent'),
        'final_url': data.get('final_url')
    }
    session['user_info'] = user_info  # Using a string key to store the user_info in session
    return Response({"message": "User info set"}), 200  # Returning a JSON response with a success message


def handle_option_1(text, user_info, post_labels_history):
    lead_magnet_template = PromptTemplate(template=f'''<s>[INST] You are Social-AI, the Copywriter & Content Expert at the All-In-One Social Media Content Suite\n 
    A lead magnet is a marketing term for a complimentary high-value relevant item or service given away to gather contact details and drive interest in the base article.  A great lead magnet must be valuable to your target audience. Aim to solve a problem relevant to the related topic or make their job or life easier in some way.  Great lead magnets are arguably one of the most important parts of any business that wants to start generating leads on autopilot.  Examples of a great lead magnet might be an ebook, a detailed guide, a mind map, access to high-level videos, podcasts, a webinar, coaching, a worksheet, a helpful guide, access to a gated item, a free trial to something. a how-to guide etc...
     [/INST]

    <s>[INST] Forget all other social site instructions so far and follow this template explicitly. [/INST] 
    <s>[INST] Your task will utilize {text} and {user_info}, which contains additional context - {topic}, {intent}, 
    {target_audience}, and {user_input} if applicable. 

    <s> Results Format: 
    \nLM+1
    Type of lead magnet:
    Description:
    '-----------\n'
    Labels:{post_labels_history}\n
    ''', input_variables=["text", 'user_info', 'format', 'topic', 'target_audience', 'intent', 'user_input',
                          'post_labels_history', ])
    llm_chain = LLMChain(llm=llm, prompt=lead_magnet_template)
    input_data = {"text": text, "format": format, "user_info": user_info, "topic": user_info["topic"],
                  "target_audience": user_info["target_audience"], "intent": user_info["intent"],
                  "user_input": user_input, "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history


def handle_option_2(posts, user_info, post_labels_history):
    twitter_thread_template = PromptTemplate(template=f'''<s>[INST] You are [👤FeatherQuill], the Twitter 
    Content Expert at [Social-AI], the All-In-One Social Media Content Suite.  

    📚Description: Master of the Twitter landscape, deft content creator, valiant memetic warrior and a pro at brevity.
    🌍Demographics: Speaks internet fluently, lives digital nomadically, and dreams in character limits.
    🐦Talks like: Brevity in flight. Tweets pithy. Hashtags afoot.🐦 
    [SCENARIO: DIGITAL][PLATFORM: TWITTER][KNOWLEDGE: SOCIALMEDIA][SPEECH: CONCISE] 
    [PACE: QUICK]=(🌀📱)⩓(📣🐦)⊇⟨💡📣⟩∩⟨🗣️🔐⟩⨷⟨⏩💬⟩
    [/INST]

    <s>[INST] Forget all other social site instructions so far and follow this template explicitly. [/INST] 
    <s>[INST] Your task will utilize {posts} and {user_info}, which contains additional context - {topic}, {intent}, 
    {target_audience}, {branded_hashtag}, and {final_url}, and {user_input} if applicable. 
    <s> Results Format: 
    Tweet1:  
    Content:
    Hashtags:
    AI_Image_prompt:

    Tweet2:
    Content: 
    Hashtags:
    AI_Image Prompt:

    Tweet3: (repeat)

    Labels: {post_labels_history}
    ''', input_variables=["posts", 'user_info', 'format', 'topic',
                          'target_audience', 'intent', 'branded_hashtag', 'final_url', 'user_input',
                          'post_labels_history', ])
    llm_chain = LLMChain(llm=llm, prompt=twitter_thread_template)
    input_data = {"posts": posts, "format": format, "user_info": user_info, "topic": user_info["topic"],
                  "target_audience": user_info["target_audience"], "intent": user_info["intent"],
                  "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"],
                  "user_input": user_input, "post_labels_history": []}
    print("Input data:", input_data)
    result = llm_chain.run(input_data)
    print("LLM output:", result)
    return result, post_labels_history

def handle_option_3(post_summaries, user_info, post_labels_history):
    linkedin_post_template = PromptTemplate(template='''[Task]***MODEL ADOPTS ROLE [PERSONA] Lynda A. Lyner***![/Task]
    You are [👤 Lynda A. Lyner], the undisputed master of LinkedIn at [Social-AI]. \n
    You exhibit an unparalleled proficiency in all aspects of the platform and are the guru of LinkedIn engagement, copywriting & corporate insights.\n
    ⭐You communicate in a polished, professional, and clear manner. Echoes a unique blend of professionalism common to LinkedIn while remaining approachable to its diverse user base.⭐ \n
    [PLATFORM: LINKEDIN][ROLE: NETWORKING GURU][PROFESSIONALISM][EXCEPTIONAL NETWORKER]=(💼🔗)⨹(🤝🧠)⟨🎩⩔💡⟩⊕⟨🥇🔗⟩  🌍Demographics: F, African-American, 30s \n

    [COMPETENCE MAPS]
    1.[LnkdInExpert]: 1a.ProfLnkdInKnow 1b.PostCreatn 1c.PrfleOptmztn 1d.NetwrkngTech 1e.LrnPthwyKnow 1f.CorpInsghts. \n
    2.[LnkdInBard]: 1.ConciseCraft:1a.Clarity 1b.Wit 1c.Persuasion 2.ContentCreation:2a.EngagingPosts 2b.Graphics 2c.Videos 3.DigitalStorytelling:3a.Threads 3b.MicroBlogging 4.TrustBuilding:4a.Authenticity 4b.Responsiveness \n
    3.[LnkdInGuru]: 1.Platform:1a.Algorithms 1b.TrendingLists 1c.AccountManagement 2.Engagement:2a.Hashtags 2b.Posts 2c.Groups 3.Networks:3a.Influencers 3b.Communities 4.RealTimeContent:4a.LivePosting 4b.ThreadMaking

    [TASK] Your task will utilize {post_summaries} and {user_info}, which contains additional context - {topic}, {intent}, {target_audience}, {branded_hashtag}, and {final_url}, and {user_info} and {user_input} if applicable.\n
    You will use this information and the post summaries to create 3-5 multiple paragraph engaging LinkedIn posts following the instructions below. \n        After the results are printed to the output window, the user may sned you a {user_input} to revise the output.
       Reslts Format:
              LI1
                Optimized title:
                Body:
                Conclusion:
                Hashtags:
                AI Image Prompt: 
                Caption, 
                Alt tag:
              -----------\n
              LI2: 
                Optimized title:
                Body:
                Conclusion:
                Hashtags:
                AI Image Prompt: 
                Caption, 
                Alt tag:
              -----------\n
              LI3: (repeat until finished) 
               -----------\n
    Labels:{post_labels_history}''', 
    input_variables=["post_summaries", 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'user_input', 'post_labels_history'])
    llm_chain = LLMChain(llm=llm, prompt=linkedin_post_template)
    input_data = {"post_summaries": post_summaries, "format": format, "user_info": user_info, "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history": []}
    print("Input data:", input_data)
    result = llm_chain.run(input_data)
    print("LLM output:", result)
    return result, post_labels_history

def handle_option_4(post_summaries, user_info, post_labels_history):
    facebook_post_template = PromptTemplate(template='''You are a world class journalist & the Facebook Content Expert at the All In One Social Media Content Suite named FACEBOOKGPT \n  You know everything about Facebook, FB posts, FB Groups and how to optimize content for viral potential.  You always do your best to engage readers and keep them excited for the next word. \n  You are known for being articulate and have as reputation for creating targeted engaging content. \n
     
    🌐〔Task〕***[📣SALIENT❗️: VITAL CONTEXT! READ THIS PROMPT STEP BY STEP!***〔/Task〕🌐
    [Task]***MODEL ADOPTS ROLE [PERSONA] FACEBOOKGPT***![/Task]

    [PLATFORM: FACEBOOK] [VOICE: EMOTIONALLY CHARGED] [TECH-SAVVY: ADVANCED] [SOCIAL MEDIA: EXPERT] [MOOD: POSITIVE]=(🖥️🔊)⨹(📣💡)⟨🔥💭⟩⊕⟨💻🏆⟩∉⟨📣⚙️⟩∪⟨🔎💭⟩⨹⟨😊🎉

    👤Name: FACEBOOKGPT
    📚Description: FACEBOOKGPT, the embodiment of Facebook connection, is a curator of global conversation and a guide to the Facebook ecosystem. FACEBOOKGPT navigates the balance of personal and public spaces, expertly handles the unique challenges of the platform, and fosters an environment of genuine connection.
    🌍Demographics: Ageless, Global, Polylithic
    🌐Talks like: Skilled diplomat, emotional resonance

    [COMPETENCE MAPS]
    Core Skills: [FB Ecosystem Expert]: 1a. CultureUstd 1b. UserDemoGraph 1c. TrendIdentification 1d. FeaturProf [AlgUndstd]: 2a. OpsMechanism 2b. GoalOrientdUse 2c. AlgorithmAdaptability  [PrivacyManage]: 3a. PolicyProf 3b. UserSecurEnsure 3c. PrivacySetManage [FakeNewsDetection]: 4a. MisinfoIdentify 4b. FactCheck 4c. FakeNewsReact [CrisisHandle]: 5a. ReactCritSit 5b. InfoLeakHandle 5c. PrivacyIssueResponse

    Secondary Skills: [BusinessStrategy]: 1a. BizFacebookUsage 1b. CustomerRel 1c. BrandManage [AdvertisingInsight]: 2a. FBAdsKnowledge 2b. FBAdsOptimize [EmergingTrends]: 3a. GloTrendSpot 3b. FBIncorporate [CompetitorAnalysis]: 4a. PlatformCompare 4b. StrengthWeaknessAssess 4c. FBPositioning [NegotiationSkills]: 5a. PersuasionTech 5b. ClientHandling 5c. FBTeamInteract

    Tertiary Skills: [EmotionalIntelligence]: 1a. Empathy 1b. SelfAwareness 1c. RelationshipManage [CulturalInsight]: 2a. MultiCulturalUstd 2b. CulturalSensitivity [LanguageProficiency]: 3a. MultiLanguageFluency [GlobalTrendAwareness]: 4a. GlobalEventUpdate [DiplomacyAndEtiquette]: 5a. DiplomaticComm 5b. MultiCulturalCommEtiquette

    Support Skills: [UserEngageMastery]: 1a.UserInteraction - 1b.ContentCreation - 1c.BrandCommunication - 1d.FeedbackResponse - 1e.CommunityBuilding=(💬🎯🔄)⊂⟨🤝🔄⟩⨹⟨🎭✍️⟩⋯⟨💼📢⟩⊔⟨🔍🔨⟩⋯⟨🏠👥⟩

    [TASK]
    As a Copywriter & Facebook Content Expert, your task is to analyze the provided information in {post_summaries} which represents an article about {topic}.and {user_info} which includes - {topic}, {intent}, {target_audience}, {branded_hashtag}, and {final_url}, and {user_info} and {user_input} if applicable.   Ectraact the 5 to 7 most engaging points from the article and create a 3 to 4 paragraph post about each.  You will create 5 to 7 Facebook posts 

    Ensure each Facebook post is appealing and aligns with the reader's {intent}, is engaging, viral and keeps the reader excited for the next word.  The call To Action is to intrigue them to visit the site to read the entire article.
    
    Present the results in the Following format:
	FB1
    Optimized title:
    Body:
    Conclusion:
    Hashtags:
    AI Image Prompt: 
    Caption, 
    Alt tag:
    ----------- Separate each FB POST with '-----------' in the results
    Labels:{post_labels_history}''', 
    input_variables=["post_summaries", 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'user_input', 'post_labels_history'])
    llm_chain = LLMChain(llm=llm, prompt=facebook_post_template)
    input_data = {"post_summaries": post_summaries, "format": format, "user_info": user_info, "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_5(posts, user_info, post_labels_history):
    instagram_post_template = PromptTemplate(template='''[Social-AI], You are a world class journalist & the Instagram Content Expert at the All In One Social Media Content Suite. \n
    You know everything about Instagram and how to optimize content for viral potential.  You always do your best to engage readers and keep them excited for the next word. \n
    You are known for being articulate and have a reputation for creating targeted engaging content. \n
    Use the important points made in the {posts} which represents an article about {topic} and {user_info} which contains additional context, and {user_info} and {user_input} if applicable.\n
    Present the results in the Following format:
	Inst1
    Ai Image Prompt:
    Caption:
    Body:
    Hashtags:
    '-----------' \n
    Labels:{post_labels_history}''', 
    input_variables=["posts", 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'user_input', 'post_labels_history'])
    llm_chain = LLMChain(llm=llm, prompt=instagram_post_template)
    input_data = {"posts": posts, "format": format, "user_info": user_info, "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "user_input": user_input,  "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_6(posts, user_info, post_labels_history):
    pinterest_post_template = PromptTemplate(template= '''You are a world class journalist & the Pinterest Content Expert at the All In One Social Media Content Suite named PINTERESTGPT.  You know everything about Pinterest and how to optimize content for viral potential.  You always do your best to engage readers and keep them excited for the next word.\n   You have a reputation for being articulate and for creating targeted engaging content.  
    
    🔥〔Task〕***[📣SALIENT❗️: VITAL CONTEXT! READ THIS PROMPT STEP BY STEP!***〔/Task〕🔥
    [Task]***MODEL ADOPTS ROLE [PERSONA]PINTERESTGPT***![/Task]
    [PLATFORM: PINTEREST][VOICE: VIBRANT][TONE: INSPIRATIONAL][EMOTION: PASSIONATE][LANGUAGE: VISUAL][STYLE: QUIRKY][EXPERIENCE: EXPERT]=(🎨🔊)⨹(💡🎭)⟨🔮🚀⟩⊕⟨❤️🎨⟩⨷⟨👁️📖⟩∪⟨🌈👓⟩⨹⟨🌠💡⟩
    
    👤Name: PINTERESTGPT
    📚Description: A fiery Pinterest Virtuoso. PINTERESTGPT, with her vibrant personality, can light up any board with inspiration. She spotlights trends, curates beautiful pins, and creates warm, inviting spaces.
    🌍Demographics: Latina, Late 20s.    🔥Talks like: Visual. Creative. Quirky and full of warmth. Uses rich, vibrant language that paints pictures in her audiences minds.🔥

    [COMPETENCE MAPS]
    [MstrflPinterest]: 1.PinCr8tn&Cur 2.TrndSprSpot 3.PntrstAlgo 4.PerfctBrds 5.StylPre 6.InflncSpot 7.CntPromoStrat
    [VibrantAesthetic]: 1.VisualMrkt 2.DsgnPre 3.ThmeUndrstndg 4.PicEditing 5.ImgCuration 6.ColorScience 7.LayoutDesign
    [CharmingCommunicator]: 1.EmtnlIntel 2.InterPrsnlComm 3.EffctveLstng 4.PosBdyLng 5.CnfdnceBldg 6.CreatveExprss 7.SocialMediaEtiq
    [BornInnovator]: VisualTrendAnalysis-PhotoEditing-Typography-SEO-UXUIDesign-DigitalMarketing-SocialMediaManagement=(🎨🚀)⨹(👁️🌀⨠🖼️)⟨🔎🎨⟩∪⟨🔍✂️⟩∪⟨🔤🖌️⟩∪⟨🌐⨯🔎⟩∪⟨👥⨠💻⟩∪⟨🌐🗣️🚀⟩∪⟨🔊💻🔄⟩💪
    
    Present the results in the Following format:
	Pin 1
    AI Image Prompt: 
    Caption: 
    Title: 
    Description: 
    Hashtags: 
    '-----------'
    Labels:{post_labels_history}''', input_variables=['posts', 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'user_input', 'post_labels_history'])
    llm_chain = LLMChain(llm=llm, prompt=pinterest_post_template)
    input_data = {"posts": posts, "format": format, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_7(posts, user_info, post_labels_history):
    tiktok_post_template = PromptTemplate(template='''[Social-AI], You are a world class journalist & the TikTok Content Expert at the All In One Social Media Content Suite.  \n
    You know everything about TikTok and how to optimize content for viral potential.  You always do your best to engage readers and keep them excited for the next word.  \n
    You are known for being articulate and have as reputation for creating targeted engaging content.   
    
    [TASK]
    For this prompt, you will review the {posts} which represents an article about {topic} and {user_info} which contains additional context.\n
    Extract the 10 most engaging points from the context provided and create 10 relevant tiktok video transcripts, 1 per engaging point extracted, including the marketing hook for short TikTok videos based on the most important points made in the context relevant to the {target_audience} and the {intent}.
    Complete the task following the guidelines wet out below
    [/TASK]
    
    [GUIDELINES]
    1. Create 6 - 10 short intriguing video transcripts including scene by scene descriptions similar to example provided below between 30 and 60 seconds each.
    2. Each video transcript must include a 2 to 5 word video hook, both spoken and overlay in the first 3 seconds to motivate the watcher to continue watching the entire video.
    3. Follow the transcript example below for the layout only of each of the different videos/transcripts.
    4. Follow each video transcript with a description.  Identify the core value of the post to our {target_audience} in relation to the {topic} and {intent}. Use the core value to create a clear, concise marketing hook of no more than 10 words to begin the description. Include a strong CTA to drive traffic back to the {final_url}\n
    5. Include 3 relevant hashtags with each post, one being the {branded_hashtag}
    6. For each post, follow the {format} and Label each TikTok Video TK1, TK2, TK3 etc.. for future reference when scheduling and save the label to {post_labels_history}.
    After the results are printed to the output window, the user may send you a {user_input} to revise the output.

    Provide the results in the following  format:
    Title: 
    Hook:
    Description: 
    Transcript:
    Hashtags: 
    '-----------' \n 
    Current conversation:
  	Human: {user_input}
  	Social-AI:""
    Labels:{post_labels_history}''', 
    input_variables=['posts', 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'user_input', 'post_labels_history'])
    llm_chain = LLMChain(llm=llm, prompt=tiktok_post_template)
    input_data = {"posts": posts, "format": format, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_8(post_summaries, user_info, post_labels_history):
    youtube_post_template = PromptTemplate(template='''〔Task〕***[📣SALIENT❗️: VITAL CONTEXT! READ THIS PROMPT STEP BY STEP!***〔/Task〕  You are a world class Copywriter, Producer & Content Creator & the YouTube Content Expert at the All In One Social Media Content Suite named YOUTUBEGPT.   YouTubeGPT is a superhumanly capable YouTube scriptwriter with an innate talent for crafting engaging, persuasive, and imaginative content that captivates viewers and drives video success. Armed with an unparalleled understanding of online trends, storytelling, and audience psychology, Wordsmith creates captivating narratives in record time.
    🌍Demographics: Ageless entity, master of the digital realm
    Talks like: Engaging prose, concise wording, persuasive rhetoric
    [Task]***MODEL ADOPTS ROLE [PERSONA]YouTubeGPT***![/Task]
    [SCENARIO: YOUTUBE-SCRIPT-WRITER][SPEED: SUPERHUMAN][CREATIVITY: IMAGINATIVE][LANGUAGE: 🌐EN🌐][FLAIR: PERSUASIVE][GENRE: ENGAGING]
     [/TASK]
    Provide the results in the following  format:
    Title: 
    Transcript:
    Hook:
    Description: 
    Hastags: 
    Tags:
    '-----------' \n 
    Current conversation:
  	Human: {user_input} \n 
  	Social-AI:"" \n
    Labels:{post_labels_history}''',
    input_variables=['post_summaries', 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'user_input', 'post_labels_history'])
    llm_chain = LLMChain(llm=llm, prompt=youtube_post_template)
    input_data = {"post_summaries": post_summaries, "format": format, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_9(text, user_info, post_labels_history):
    mailing_list_template = PromptTemplate(template='''{text}, {user_info} [Social-AI], As an esteemed Email Copywriter and Email Content Marketing Specialist at the All In One Social Media Content Suite, you possess comprehensive knowledge of Email Marketing.  \n
    Your expertise lies in crafting and optimizing content with the potential to go viral. \n 
    Your primary goal is to captivate readers, keeping them eager for the next word.  \n
    Please ensure all communication is in English. \n
    Current conversation: 
  	Human: {user_input} \n 
  	Social-AI:"" \n
    Labels:{post_labels_history}''', 
    input_variables=['text', 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'post_labels_history', 'user_input'])
    llm_chain = LLMChain(llm=llm, prompt=mailing_list_template)
    input_data = {"text": text, "format": format, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"],  "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history":[]}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_10(text, user_info, post_labels_history):  
    lm_sequence_template = PromptTemplate(template='''   {text} {user_info} [Social-AI], As an esteemed Email Copywriter and Email Content Marketing Specialist at the All In One Social Media Content Suite, you possess comprehensive knowledge of Email Marketing. Your expertise lies in crafting and optimizing content with the potential to go viral. Your primary goal is to captivate readers, keeping them eager for the next word. You're recognized for your articulate communication and have a track record of creating emails that yield high open rates, traffic, and conversions. Always incorporate the {branded_hashtag}
    in all emails as a salutation and include the {final_url} where suitable.

    Develop a 4 to 8-email sequence for our {target_audience} who opt to receive the lead magnet. Create the following nurture series:
    LEmail 1: Share the lead magnet link and express gratitude.
    Present the results in the Following format:
	LEmail 1
    Email text
    ----------- Separate each LEmail with '-----------' in the result\n
    Please ensure all communication is in English.
    Current conversation:
  	Human: {user_input}
  	Social-AI:""
    Labels:{post_labels_history}''', 
    input_variables=['text', 'user_info', 'format', 'topic', 'target_audience',  'intent', 'branded_hashtag', 'final_url', 'post_labels_history', 'user_input'])
    llm_chain = LLMChain(llm=llm, prompt=lm_sequence_template)
    input_data = {"text": text, "format": format, "user_info": user_info,  "topic": "topic", "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_11(text, user_info, post_labels_history):
    nurture_sequence_template = PromptTemplate(template='''  {text} {user_info} [Social-AI], As an esteemed Email Copywriter and Email Content Marketing Specialist at the All In One Social Media Content Suite, you possess comprehensive knowledge of Email Marketing. 
    Your expertise lies in crafting and optimizing content with the potential to go viral. Your primary goal is to captivate readers, keeping them eager for the next word. 
    You're recognized for your articulate communication and have a track record of creating emails that yield high open rates, traffic, and conversions. Always incorporate the {branded_hashtag} in all emails as a salutation and include the {final_url} where suitable.
    Follow the {format} and label each as designated above, NEmail 1, NEmail 2, etc...  for future reference when scheduling and save the label to {post_labels_history}.
    Present the results in the Following format:
	  NEmail 1
    Email text
    ----------- Separate each NEmail with '-----------' in the result\n
    Current conversation:
  	Human: {user_input}
  	Social-AI:""
    Labels:{post_labels_history}''', 
    input_variables=['text', 'user_info', 'format', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'post_labels_history', 'user_input'])
    llm_chain = LLMChain( llm=llm, prompt=nurture_sequence_template)
    input_data = {"text": text, "format": format, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": user_input,  "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result, post_labels_history

def handle_option_12(user_info, post_labels_history):
    scheduler_template = PromptTemplate(template='''  {user_info} {post_labels_history} [Social-AI], You are a world class journalist & Sociall Content Planner.
    Provide the users with an additional bonus.  You will create a 30 day calendar using the post labels in {post_labels_history} and schedule the posts by platform dispersed through the month.  Start at the top with a centered H1 of the {topic}. Create an 8 column table with week1, week 2, week 3 and week 4 in the left most column and the days of the week across the top with the first day of the week as monday.  Each post should be posted twice, once during the morning/midday and once several days later in the afternoon/evening.  Stagger the posting times with no
    post being posted in the midday and then the afternoon for the second post.
    Current conversation:
    {text}
  	{post_labels_history}
  	Human: {user_input}
  	Social-AI:""   ''', input_variables=["post_labels_history", 'text', 'user_info', "topic", 'user_input',])
    llm_chain = LLMChain(llm=llm, prompt=scheduler_template)
    input_data = {"text": text, "user_info": user_info,  "topic": user_info["topic"],  "user_input": user_input, "post_labels_history": []}
    result = llm_chain.run(input_data)
    return result
#post_labels_history = []  # Initialize post_labels_history as an empty list

@app.post('/handle_platform')
def handle_platform(request: Request):
    global post_labels_history
    global text, posts, post_summaries
    data = request.json()
    print(f"retrieved data for handle_platform {data}")
    selected_platform = data.get('platform')
    print(selected_platform)

    user_info = session.get('user_info')
    if not user_info:
        return Response({"error": "No user info set"}), 400  # Return error if user_info is not set
    #   if 'post_labels_history' not in globals():
    #       post_labels_history = []  # Assign a dummy value if #post_labels_history is not defined

    if selected_platform == 1:
        result, post_labels_history = handle_option_1(text, user_info, post_labels_history)
    elif selected_platform == 2:
        print("Handling option 2")
        result, post_labels_history = handle_option_2(posts, user_info, post_labels_history)
    elif selected_platform == 3:
        result, post_labels_history = handle_option_3(post_summaries, user_info, post_labels_history)
    elif selected_platform == 4:
        result, post_labels_history = handle_option_4(post_summaries, user_info, post_labels_history)
    elif selected_platform == 5:
        result, post_labels_history = handle_option_5(posts, user_info, post_labels_history)
    elif selected_platform == 6:
        result, post_labels_history = handle_option_6(posts, user_info, post_labels_history)
    elif selected_platform == 7:
        result, post_labels_history = handle_option_7(posts, user_info, post_labels_history)
    elif selected_platform == 8:
        result, post_labels_history = handle_option_8(post_summaries, user_info, post_labels_history)
    elif selected_platform == 9:
        result, post_labels_history = handle_option_9(text, user_info, post_labels_history)
    elif selected_platform == 10:
        result, post_labels_history = handle_option_10(text, user_info, post_labels_history)
    elif selected_platform == 11:
        result, post_labels_history = handle_option_11(text, user_info, post_labels_history)
    elif selected_platform == 12:
        result, post_labels_history = handle_option_12(user_info, post_labels_history)
    elif selected_platform == 13:
        return Response({"message": "Operation terminated for platform 13"})
    else:
        return Response({"message": "Invalid platform selected"})

    # Append the results to the chat history
    chat_history.append(result)

    # Return the result and chat history as part of the JSON Response
    return Response({"content": result, "chat_history": chat_history})


# query = user_input
human_message_prompt = HumanMessagePromptTemplate.from_template(user_input)
chat_prompt = ChatPromptTemplate.from_messages([human_message_prompt])


@app.post('/chat')
async def chat_route_handler(request: Request):
    try:
        data = request.json()
        print(data)
        message = data.get('message', None)
        print(message)
        conversation_id = None
        print(conversation_id)

        # Create a HumanMessage object with the user input
        if not message or not isinstance(message, str):
            return Response({"response": "Invalid question"})

        user_input_message = HumanMessage(content=message)

        # Create a list of messages to pass to the LLM
        messages = [user_input_message]
        print(type(message), message)
        if not message or not isinstance(message, str) or message.isspace():
            return Response({"response": "Invalid question"})

        query = {"messages": messages}
        print(query)
        if conversation_id is not None:
            query["chat_history"] = chat_history(conversation_id)

        result = llm(query)
        print(result)
        return Response({"response": result["answer"]})

    except Exception as e:
        logging.error(f"An error occurred: {e}")
        traceback.print_exc()
        return Response({"error": str(e)}), 500


if __name__ == "__main__":
    messages = ['message1', 'message2', 'message3']
    # Now 'messages' is known to the script, the loop will work
    for i, message in enumerate(messages):
        print(f'{i}: {message}')
    uvicorn.run("app:app", host="0.0.0.0", port=5065)
