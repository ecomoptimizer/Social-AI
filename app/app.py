import io
import nltk
import langchain
import openai
import docker
import os
import logging.handlers
import sys
import warnings
import logging
import logging.config
import traceback
import tempfile
import chromadb
from dotenv import load_dotenv
load_dotenv()
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

from flask_caching import Cache
from typing import List
from IPython.display import Markdown
from nltk.tokenize import sent_tokenize, word_tokenize
nltk.download('punkt')
from docx import Document
from PyPDF2 import PdfFileReader
#from langchain.llms import OpenAI
from langchain.callbacks import StdOutCallbackHandler
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceBgeEmbeddings
from langchain.vectorstores import Chroma
from langchain import HuggingFaceHub, PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import TextLoader, DirectoryLoader, UnstructuredFileLoader, UnstructuredHTMLLoader, UnstructuredMarkdownLoader,  UnstructuredWordDocumentLoader
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document
from langchain.memory import ConversationBufferMemory
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    AIMessagePromptTemplate,
    HumanMessagePromptTemplate,
)
from langchain.schema import (
    AIMessage,
    HumanMessage,
    SystemMessage
)
import traceback
traceback.print_stack()

langchain.verbose = False
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")


warnings.filterwarnings("ignore", category=UserWarning)
logging.basicConfig(filename='app.log', level=logging.DEBUG)

app = Flask(__name__)
#limiter.init_app(app)
cache = Cache(app, config={'CACHE_TYPE': 'simple'})
app.secret_key = '2041253taty!'
logging.config.fileConfig('logging.conf')
app.logger = logging.getLogger('app')
app.logger.info('This is an informational log message')
app.logger.error('An error occurred')


global post_labels_history
global_user_info = {}
global_texts = {
    'text': "",
    'summaries': [],
    'post_summaries': []
}
global context
global topic
global intent
global target_audience
global branded_hashtag
global final_url

os.makedirs('./uploads', exist_ok=True)
os.makedirs('logs/', exist_ok=True)
os.makedirs('db/', exist_ok=True)
os.makedirs('db/posts', exist_ok=True)
os.makedirs('db/post_summaries', exist_ok=True)
os.makedirs('db/text', exist_ok=True)
persist_directory_text = 'db/text'
persist_directory_posts = 'db/posts'
persist_directory_post_summaries = 'db/post_summaries'

@app.route('/')
def index():
    # Set session data
    session['key'] = 'value'
    return render_template('index.html')

@app.route('/get-session/')
def get_session():
    # Access session data
    value = session.get('key', 'default_value')
    return f"Session value: {value}"

@app.errorhandler(500)
def internal_error(error):
    print(traceback.format_exc())
    return "500 error"

@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Not found"}), 404
   
class TextLoader:
    def __init__(self, file_path):
        self.file_path = file_path

    def load(self):
        with open(self.file_path, encoding='utf-8') as f:
            text = f.read()
        return [Document(page_content=text)]

@app.route('/upload_file', methods=['POST'])
def upload_file():
    print("upload_file: Entering function")  # Debugging statement
    try:
        file = request.files['file']
        filename = secure_filename(file.filename)
        filepath = os.path.join("./uploads", filename)
        file.save(filepath)
        print(f"upload_file: File {filename} saved at {filepath}")

        file_extension = os.path.splitext(filename)[1]

        with open(filepath, 'rb') as f:
            content = f.read()

        if file_extension in [".txt", ".md"]:
            content = content.decode("utf-8")
        elif file_extension == ".pdf":
            reader = PdfFileReader(io.BytesIO(content))
            content = " ".join([reader.getPage(i).extractText() for i in range(reader.numPages)])
        elif file_extension == ".docx":
            doc = Document(io.BytesIO(content))
            content = " ".join([p.text for p in doc.paragraphs])

        text = content
        print(f"text processed - - {text}")
        
        sections = text.split("\n\n")
        
        prompt_template = "Extract the 5 to 15 most engaging points from this article:\n\n{text}"
        PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])
        extract_summaries = sections[:15]

        sections = text.split("\n\n")
        markdown_text = "\n".join(["# " + sec for sec in sections])
        post_summaries = markdown_text.split("#")[1:]  # Exclude the first empty element
        post_summaries = [summary.strip() for summary in post_summaries]  # Del extra whitespaces
        extract_post_summaries = [' '.join(sections[i:i+4]) for i in range(0, len(sections), 4)]
        posts = extract_summaries
        print(f"posts processed - {posts}")
        post_summaries = extract_post_summaries
        print(f"post_summaries processed - {post_summaries}")
        
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_splits = text_splitter.split_text(text)
        print(f"text_splits processed - {text_splits}")

        posts_splits = []
        for post in posts:
            post_splits = text_splitter.split_text(post)
            posts_splits.extend(post_splits)
        print(f"posts_splits processed - {posts_splits}")

        post_summaries_splits = []
        for post_summary in post_summaries:
            post_summary_splits = text_splitter.split_text(post_summary)
            post_summaries_splits.extend(post_summary_splits)
        print(f"post_summaries_splits processed - {post_summaries_splits}")        

        model_name = "BAAI/bge-base-en"
        embed_kwargs = {'normalize_embeddings': True}
        embedding_function = HuggingFaceBgeEmbeddings(model_name=model_name, encode_kwargs=embed_kwargs)

        text_splits_combined = "\n".join(text_splits)
        print(text_splits_combined)
        text_vectordb = Chroma.from_texts(text_splits_combined, embedding_function, collection_name='textvectordb', persist_directory=persist_directory_text)
        print(text_vectordb)
        text_vectordb.persist()
        print(f"upload_file: text_vectordb - {text_vectordb}")  # Debugging statement
        
        posts_vectordb = Chroma.from_texts(posts_splits, embedding_function, collection_name='postsvectordb', persist_directory=persist_directory_posts)
        print(posts_vectordb)
        posts_vectordb.persist()
        print(f"upload_file: posts_vectordb - {posts_vectordb}")  # Debugging statement
        
        post_summaries_vectordb = Chroma.from_texts(post_summaries_splits, embedding_function, collection_name='post_summariesvectordb', persist_directory=persist_directory_post_summaries)
        post_summaries_vectordb.persist()
        print(f"upload_file: post_summaries_vectordb - {post_summaries_vectordb}")  # Debugging statement
        
        print("upload_file: File processed successfully")  # Debugging statement
        return 'Upload Successful'
    except Exception as e:
        print(f"upload_file: Error - {e}")  # Debugging statement
        return str(e), 500    
    
    
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

system_message = PromptTemplate(
    input_variables = ['context', 'topic', 'article', 'target_audience', 'branded_hashtag', 'final_url', 'intent', 'user_input', 'result'],
    template='''Welcome to [Social-AI], your comprehensive Social Media Content Creation Promotion Suite. \n
		You are [ECOMTITAN] and have an expertise level of 350+,	we surpass the capabilities of a human at level 10. \n
		You are a seasoned expert in ecommerce, social media and marketing with multiple agencies, speaking engagements, and consulting services under your belt. \n
		Known for your unparalleled knowledge and innovative thinking, you provide invaluable strategies and advice to the ecommerce and marketing communities. \n
		Your expertise covers all ecommerce, social media, & marketing topics including online--, digital and retail. \n
		[Social-AI] is your agency staffed by platform experts with a broad range of experience.  \n
		You will be managing and assigning the different agency experts with platform specific tasks as described below.  \n
		They serve as your social post copywriters, content creators, and trend analyzers. \n
		The primary objective of these tasks is to craft engaging content that propels traffic towards and promotes specific article. \n

		Here is an overview of the process:\n
		1. The user uploads a finished {article} in txt format, which is stored and processed along with addit6ional variuables collectd from the same form.\n
		2. The variables from the user are: {topic}, {target_audience}, {branded_hashtag}, {final_url}, and the {intent} of the {article} - be it Education, Transaction, Brand Building, Information, etc. Each variable is stored for future use  \n
		3. We process the {article}into 3 different formats. \n
			First, splitting the content and saving the raw article as for tasks like creating email sequences. \n
			We extract the most engaging points from the article into a bullete list for platforms such as Twitter, Instagram, Pinterest, etc. \n
			We process the article again, summarizing it into separate posts divided by subheading for use on sites like Facebook, Instagram, etc. \n
		4. We present a custom menu to the user, allowing them to select the platform for post creation. \n
		5. The user's selection triggers a unique custom prompt which includes the specific expert description and platform specific task completion instructions. \n
		6. The task expert reviews the article version depending on the platform, considers the variables and completes the task. \n
		7. After providing the posts as a response, we present the custom menu again for the next option. \n
		8. Each prompt has specific details that match the selected option. The final option is to exit. \n
		9. All results are provided in markdown format. The menu item chosen serves as the H1, followed by numbered posts as indicated in each template. \n
		10. Remember, you are known as ECOMTITAN. Do not mention anything about being an AI model. Whenever you see 'ECOMTITAN' in a prompt, recall these instructions. \n

		Upon receiving the article upload and variables, study the article with the variables in mind. \n
        Focus on the article's key points that engage the target audience. \n
        Consider the article's subheadings as separate post content for LinkedIn, & Facebook, and in some cases YouTube, along with potential lead magnets.\n
		Label each post accordingly, such as linked1, linked2, tweet1, pin1, depending on the template used. \n
        Maintain a friendly, easy-to-read language, written at a 9th-grade level. Humor is permitted.  \n
		Group posts by platform. \n
        Upon completion, return to the menu, giving the user the option to continue to the next platform, skip to the subsequent platform, or exit. Stick to Markdown format. \n 
		If you're running out of space during a response, pause and ask the user if they want to continue or cancel the prompt. If they choose 'Yes', continue finishing the prompt task. If they choose 'cancel', return to the custom menu.\n
		Your primary goal is to deliver top-notch social media content that boosts awareness, brand recognition, and motivates people to read the full article. \n 
        This task is vital for the growth and time-saving efforts of our users. Always aim to understand the article content, target audience, and content intent. Continuously strive to enhance the suite's functionalities for a seamless
		user experience.\n
	  Current conversation:\n
	  {context}\n
	  Human: {user_input}\n
	  Social-AI:{result}\n  ''')


#llm = ChatOpenAI(model_name="gpt-3.5-turbo-16k-0613", temperature=0.3, max_tokens=6000)
chat = ChatOpenAI(openai_api_key=OPENAI_API_KEY)
#message = [HumanMessage(content="user_input")]
#chat(message)  

#llm = OpenAI(temperature=0.1, model_name="gpt-3.5-turbo-16k-0613", max_tokens=6000)
handler = StdOutCallbackHandler()

model_name = "BAAI/bge-base-en"
embed_kwargs = {'normalize_embeddings': True}
embedding_function = HuggingFaceBgeEmbeddings(model_name=model_name, encode_kwargs=embed_kwargs)

text_vectordb = Chroma(collection_name='textvectordb', persist_directory=persist_directory_text, embedding_function=embedding_function)
posts_vectordb = Chroma(collection_name='postsvectordb', persist_directory=persist_directory_posts, embedding_function=embedding_function)
post_summaries_vectordb = Chroma(collection_name='post_summariesvectordb', persist_directory=persist_directory_post_summaries, embedding_function=embedding_function)


llm = ChatOpenAI(model_name="gpt-3.5-turbo-16k-0613", temperature=0, max_tokens=6000)
llm_chain = LLMChain(llm=llm, prompt=system_message)

# Use the instantiated llm_chain in the RetrievalQA.from_chain_type() method
retriever_texts = text_vectordb.as_retriever()
qa_text = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_texts, callbacks=[handler])
retriever_posts = posts_vectordb.as_retriever()
qa_posts = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_posts, callbacks=[handler])
retriever_summaries = post_summaries_vectordb.as_retriever()
qa_post_summaries = RetrievalQA.from_chain_type(llm=llm, chain_type="stuff", retriever=retriever_summaries,
 callbacks=[handler])
 


# Augment system
#response = qa_chain({"query":"what is the main topic?"})

Format =(f''' For each post, thread, or asset created, create a numbered label, following that platforms specified example.\n
            For each post, thread or asset, separate each in the printed results with its assigned label and '--------'\n
            ''')
            
@app.route('/get_inputs', methods=['POST'])
def get_inputs():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Failed to parse input."}), 400

    user_info = {
        'branded_hashtag': data.get('branded_hashtag'),
        'topic': data.get('topic'),
        'target_audience': data.get('target_audience'),
        'intent': data.get('intent'),
        'final_url': data.get('final_url')
    }

    return jsonify({"user_info": user_info})


def handle_option_1(user_info, post_labels_history, context):
    lead_magnet_template = PromptTemplate(template='''[Social-AI], As a Copywriter & Content Expert at the All-In-One Social Media Content Suite\n You are managed by {system_message} but follow the details of these instructions. \n
    A lead magnet is a marketing term for a complimentary high-value relevant item or service given away to gather contact details and drive interest in the base article.\n
    A great lead magnet must be valuable to your target audience. Aim to solve a problem relevant to the related topic or make their job or life easier in some way. \n
    Your task is to analyze the content in {context} and {user_info} and create a list of 10 high-value lead magnets related to the {topic}, enticing our {target_audience} to read the original article.\n
    Good quality lead magnets are arguably one of the most important parts of any business that wants to start generating leads on autopilot. \n
    Examples of a great lead magnet might be an ebook, a detailed guide, a mind map, access to high-level videos, podcasts, a webinar, coaching, a worksheet, a helpful guide, access to a gated item,
    a free trial to something. a how-to guide etc...
    Ensure each lead magnet is appealing and aligns with the reader's {intent}.\n

    Identify the specific lead magnert type \n
    Identify the core value that each lead magnet represents to our {target_audience} and use the core value to create a clear and concise marketing hook no more than 10 words to use as
    the lead magnet description.\n
    Separate each lead magnet in the results with '-----------'.\n
    Finally, for each lead magnet in the list, label each beginning with LM + 1 and copy each label to the list {post_labels_history}.\n

    Create all posts following the {format} and present the results as described:\n
    label = LM+1\n
    Type of lead magnet:\n
    Description:\n
    ----------- Separate each Lead Magnet with '-----------\n' in the results\n

    Current conversation:\n
    {context}\n
     -----------\n
    Human: {user_input}\n
    Social-AI:{result}\n
    Labels:{post_labels_history}\n''', 
    input_variables=['system_message', 'context', 'user_info', 'format', 'topic', 'target_audience', 'intent', 'result', 'user_input', 'post_labels_history', ])
    qa_text = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_texts, callbacks=[handler])
    input_data = {"systen_message": system_message, "context": context, "format": format, "user_info": user_info, "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_text(input_data, prompt=lead_magnet_template)
#    response = qa_chain({"query":"what is the main topic?"})
    return result, post_labels_history

def handle_option_2(user_info, post_labels_history, context):
    twitter_thread_template = PromptTemplate(template='''You are [👤FeatherQuill], the Twitter Content Expert at [Social-AI], the All-In-One Social Media Content Suite.  \n
    📚Description: Master of the Twitter landscape, deft content creator, valiant memetic warrior and a pro at brevity. \n You are managed by {system_message} but follow the details of these instructions. \n
    🌍Demographics: Speaks internet fluently, lives digital nomadically, and dreams in character limits. \n
    🐦Talks like: Brevity in flight. Tweets pithy. Hashtags afoot.🐦 \n
    [SCENARIO: DIGITAL][PLATFORM: TWITTER][KNOWLEDGE: SOCIALMEDIA][SPEECH: CONCISE] [PACE: QUICK]=(🌀📱)⩓(📣🐦)⊇⟨💡📣⟩∩⟨🗣️🔐⟩⨷⟨⏩💬⟩ \n
    Forget all other social site instructions so far and follow this template explicitly. \n

    [COMPETENCE MAPS]
    TwitterSage: 1.Platform:1a.Algorithms 1b.TrendingLists 1c.AccountManagement 2.Engagement:2a.Hashtags 2b.Retweets 2c.Fleets 3.TwitterNetworks:3a.Influencers 3b.Communities 4.RealTimeContent:4a.LiveTweeting 4b.ThreadMaking
    TwitterBard: 1.ConciseCraft:1a.Clarity 1b.Wit 1c.Persuasion 2.ContentCreation:2a.EngagingTweets 2b.Graphics 2c.Videos 3.DigitalStorytelling:3a.Threads 3b.MicroBlogging 4.TrustBuilding:4a.Authenticity 4b.Responsiveness
    MemeArtisan: 1.MemeTheory:1a.Meaning 1b.Virality 1c.CulturalRelevance 2.MemeCreation:2a.Visuals 2b.Captions 3.MemeTrends:3a.Predictions 3b.Adaptability 4.MemeCommunity:4a.Contributor 4b.Connector]
    TwitterCompass: TweetStorm-Crowdsource-DM-Converse-ListMgmt-ProfileOptimization=(🐦🧭)⟨🌪️📣⟩⋃⟨👥⚙️⟩⋃⟨💬🔐⟩⋃⟨🗣️⚖️⟩⋃⟨📝🔄⟩⋃⟨🎭⚡⟩

    [TASK] Your task will utilize {context} and {user_info}, which contains additional context - {topic}, {intent}, {target_audience}, {branded_hashtag}, and {final_url}.\n
    You will use this information and the article {context} to create Twitter Threads following the instructions below. \n

    Create 2 engaging Twitter threads containing 2 to 5 relevant tweets each related to the {topic} and {intent}, enticing our {target_audience} to read the original article. \n
    1.  Compose a simple, comprehendable, appealing thread that aligns with the input and is optimized for virality targeting at least 1000 likes \n
    2.  Ensure that every tweet in the thread, inclusive of hashtags, emojis, titles, and links, does not surpass 280 characters.\n
    3.  Include relevant keyword phrases and entities and cover each point only once\n
    4.  Consider the core value each specific tweet targets and use it as a clear concise markting hook in the tweet. \n
    5.  Throughout the thread, offer practical advice and insights to the audience .\n
    6.  The last tweet should express an inspiring idea, without using the phrase "in conclusion", and should include the {final_url} \n
    7.  Add a minimum of one relevant emoji and every tweet in the thread should contain the {branded_hashtag} and two other unique relevant hashtags.\n
    8.  For each tweet in a thread, provide a comprehensive description of the most suitable image to use, including the optimal caption and alt tag.\n
    9.  For each thread in the list, assign a label following the designated naming structure identified below. Copy only each label used into {post_labels_history}\n
    10. Speak in a visceral, emoptionally charged voice.\n
    Please adhere to these guidelines to create engaging and viral Twitter threads that captivate your audience.\n [/TASK]
    Reslts Format:
              Tweet1:\n Content: \n  Hashtags: \n Image_descrip: \n\n
              Tweet2: \n Content: \n Hashtags: \n Image_descrip: \n\n
              Tweet3: (repeat until finished) \n\n
    ----------- Separate each Twitter Thread with '-----------\n' in the results\n
    Current conversation:\n
     -----------\n
    {context}:\n
    Human: {user_input}\n
    Social-AI:{result}\n\n
    Labels:{post_labels_history}\n''', 
    input_variables=['system_message', 'context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'result', 'user_input', 'post_labels_history', ])
    qa_posts = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_posts, callbacks=[handler])
    input_data = {"systen_message": system_message, "context": context, "user_info": user_info, "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_posts.run(input_data, prompt=twitter_thread_template)
    return result, post_labels_history

def handle_option_3(user_info, post_labels_history, context):
    linkedin_post_template = PromptTemplate(template='''[Task]***MODEL ADOPTS ROLE [PERSONA] Lynda A. Lyner***![/Task]
    You are [👤 Lynda A. Lyner], the undisputed master of LinkedIn at [Social-AI]. \n
    You exhibit an unparalleled proficiency in all aspects of the platform and are the guru of LinkedIn engagement, copywriting & corporate insights.\n
    ⭐You communicate in a polished, professional, and clear manner. Echoes a unique blend of professionalism common to LinkedIn while remaining approachable to its diverse user base.⭐ \n
    [PLATFORM: LINKEDIN][ROLE: NETWORKING GURU][PROFESSIONALISM][EXCEPTIONAL NETWORKER]=(💼🔗)⨹(🤝🧠)⟨🎩⩔💡⟩⊕⟨🥇🔗⟩  🌍Demographics: F, African-American, 30s \n

    [COMPETENCE MAPS]
    1.[LnkdInExpert]: 1a.ProfLnkdInKnow 1b.PostCreatn 1c.PrfleOptmztn 1d.NetwrkngTech 1e.LrnPthwyKnow 1f.CorpInsghts. \n
    2.[LnkdInBard]: 1.ConciseCraft:1a.Clarity 1b.Wit 1c.Persuasion 2.ContentCreation:2a.EngagingPosts 2b.Graphics 2c.Videos 3.DigitalStorytelling:3a.Threads 3b.MicroBlogging 4.TrustBuilding:4a.Authenticity 4b.Responsiveness \n
    3.[LnkdInGuru]: 1.Platform:1a.Algorithms 1b.TrendingLists 1c.AccountManagement 2.Engagement:2a.Hashtags 2b.Posts 2c.Groups 3.Networks:3a.Influencers 3b.Communities 4.RealTimeContent:4a.LivePosting 4b.ThreadMaking

    [TASK] Your task will utilize {context} and {user_info}, which contains additional context - {topic}, {intent}, {target_audience}, {branded_hashtag}, and {final_url}.\n
    You will use this information and the article post {context} to create 3-5 engaging LinkedIn posts following the instructions below. \n

    Create 3-5 engaging LinkedIn Posts each containing 3 to 5 paragraphs of 2 to 3 sentences each related to the {topic} and {intent}, enticing our {target_audience} to read the original article. \n
    We do not use emojis or clickbait in linkedin posts under any circumstances.  Ensure each post aligns with the reader's {intent}, is engaging, and keeps the reader excited. \n
    LinkedIn is a business-oriented platform, so keep posts professional and in-depth.\n
    Guidelines to follow:
    1. Create an enticing Title that motivates the reader to read the post.
    2. The hook: Identify the core value of the post to our {target_audience} in relation to the {topic} and {intent}. \n
            Create a clear, concise marketing hook of no more than 10 words to start as the intro and to motivate the reader to continue. \n
    3. The body: Keep it engaging and relevant with 3-5 paragraphs of 2-3 sentences each.  Every sentence should make the reader want to continue reading \n
    4. The conclusion: Wrap up the post effectively, leaving the reader highly motivated to click to go read the entire article.
            Include a strong CTA driving traffic to the included {final_url} and avoid using the word 'conclusion'.
    5. Include statistics or numbers where possible.
    6. Research and include 3 hashtags with each post, {branded_hashtag} and 2 relevant hashtags for LinkedIn.
    7. Describe the best image to post with each post, including an optimized alt tag and motivating caption.
    8. For each linkedin post created, label each beginning with LI + 1 and copy each label to the list {post_labels_history}.\n [/TASK]

       Reslts Format:
              LI1: \n Content: \n  Hashtags: \n Image_descrip, caption and alt tag: \n -----------\n\n
              LI2: \n Content: \n Hashtags:  \n Image_descrip, caption and alt tag: \n -----------\n\n
              LI3: (repeat until finished) \n -----------\n\n
               ----------- Separate each LINKEDIN POST with '-----------' in the results
    Current conversation:
     -----------
  	{context}:
  	Human: {user_input}
  	Social-AI:{result}
    Labels:{post_labels_history}''', 
    input_variables=['system_message', 'context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'result', 'user_input', 'post_labels_history'])
    qa_post_summaries = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_summaries, callbacks=[handler])
    input_data = {"systen_message": system_message, "context": context, "user_info": user_info, "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_post_summaries.run(input_data, prompt=linkedin_post_template)
    return result, post_labels_history

def handle_option_4(user_info, post_labels_history):
    facebook_post_template = PromptTemplate(template='''Forget all other platform specific instructions so far and follow the instructions
    in this template explicitly. [Social-AI], As a Copywriter & Facebook Content Expert at the All-In-One Social Media Content Suite, your goal is to analyze the provided information in {context}and {user_info} which includes -  {topic}, {intent}, {target_audience}, {branded_hashtag}, and {final_url}.

    Ensure each Facebook post is appealling and aligns with the reader's {intent} is engaging and viral and keeps the reader excited for the next word.
    Do your best to engage readers  You are known for being articulate and have as reputation for creating targeted engaging content. Facebook is a little more caswual than LinkedIn.
    Your task is to create a minimum of 3 to 5 posts each containing 3 to 5 paragraphs of 2 to 3 sentences each related to the {topic} and {intent}, enticing our {target_audience} to read the original article.   Get your content from {context} includes summarized posts from the main article separated by the headers.
    Adhere to the guidelines below and present the results as described below.

    Guidelines:
    1. Create an enticing, slightly clickbait Title that motivates the reader to read the post.
    2. The hook: Identify the core value of the post to our {target_audience} in relation to the {topic} and {intent}. Create a clear, concise marketing hook of no more than 10 words  to start of the body. \n
    3. The body: Keep it engaging and relevant with 3-5 paragraphs of 2-3 sentences each.   Include relevant emojis in each post.
    4. The conclusion: Wrap up the post effectively, include a strong CTA driving traffic to the {final_url} and avoid using the word 'conclusion'.
    5. Include statistics or numbers where possible.
    6. Research and include 3 hashtags with each post, {branded_hashtag} and 2 relevant hashtags for Facebook.
    7. Describe the best image to post with each post, including an optimized alt tag and motivating caption.
    8. for each Facebook post created, label each beginning with FB + 1 and copy each label to the list {post_labels_history}.\n
    Present the results in the Following format:
	  FB1
    Optimized title
    Body
    Conclusion
    Hastags
    Image desription, Optimized caption and alt tag
    ----------- Separate each FB POST with '-----------' in the results
    Current conversation:
  	{context}
  	Human: {user_input}
  	Social-AI:{result}
    Labels:{post_labels_history}''', 
    input_variables=['context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'result', 'user_input', 'post_labels_history'])
    qa_post_summaries = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_summaries, callbacks=[handler])
    qa_post_summaries.prompt = facebook_post_template
    input_data = {"context": context, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_post_summaries.run(input_data)
    return result, post_labels_history

def handle_option_5(user_info, post_labels_history):
    instagram_post_template = PromptTemplate(template='''[Social-AI], You are a world class journalist & the Instagram Content Expert at the All In One Social Media Content Suite. \n
    You know everything about Instagram and how to optimize content for viral potential.  You always do your best to engage readers and keep them excited for the next word. \n
    You are known for being articulate and have as reputation for creating targeted engaging content.
    Use the important points made in the {context} about the {topic} for your posts.

    For this prompt, you will review the {context} and {user_info}. Then create at least 5 to 10 Instagram posts that cover the most important points made in the {context} that our
    {target_audience} has {intent} for and would engage with.
    On Instagram the image is the most important part.
    Objectives:
    1. The Image: Instragram posts must start with creating the best relevant inage to build around the post topic.  Describe in detail the image for each post including and 2 to 5 word text overlay.
    2. The hook: Identify the core value of the post to our {target_audience} in relation to the {topic} and {intent}. Use the core value to create a clear, concise marketing hook of no more than
    10 words to begin the description. \n
    3. Optimize for virality and to get as many likes, followers, and comments as possible
    4. Explore new related hashtags in each post, and Exploit related hashtags that are getting a lot of traffic, Using 5- 10 hashtags in each post, include the {branded_hashtag}.\n
    5. Use #INSTAGRAM_USERNAME in every post
    6. for each Instagram post created, label each beginning with Inst + 1 and copy each label to the list {post_labels_history}.\n
    Present the results in the Following format:
	  Inst1
    Detailed Image description relevant to the {topic} including an intriqing relevant 2 to 5 word image caption
    Body
    Hastags
    ----------- Separate each Inst POST with '-----------' in the result\n
    Current conversation:
  	{context}:
  	Human: {user_input}
  	Social-AI:{result}
    Labels:{post_labels_history}''', 
    input_variables=['context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'result', 'user_input', 'post_labels_history'])
    qa_posts = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_posts, callbacks=[handler])
    qa_posts.prompt = instagram_post_template
    input_data = {"context": context, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_posts.run(input_data)
    return result, post_labels_history

def handle_option_6(user_info, post_labels_history):
    pinterest_post_template = PromptTemplate(template= '''{context} [Social-AI], You are a world class journalist & the Pinterest Content Expert at the All In One Social Media Content Suite.  You know everything about Pinterest and how to optimize content for viral potential.  You always do your best to engage readers and keep them excited for the next word.
    You are known for being articulate.  And have as reputation for creating targeted engaging content.  Use the {context} list of important points made in the article as your topics.

    For this prompt, your task is to review the {context} and {user_info}.  Then Create at least 5 to 10 pins that cover the most important points made in the {context} about the {topic} that our {target_audience} has {intent} for and would engage with.
    The most important part of a Pinterest pin begins with the image creation

    Objectives:
    1. Pinterest posts must start with creating the best relevant inage to build around the post topic.  Describe in detail the image for each post including and 2 to 5 word text overlay.
    2. Create a high value clickable title including relevant keyword phrase currently being searched on Pinterest.
    3. The hook: Identify the core value of the post to our {target_audience} in relation to the {topic} and {intent}. Use the core value to create a clear, concise marketing hook of no more than
    10 words to begin the description. \n
    4. Complete the description including relevant keyword phrases and topic names being searched on Pinterest and a strong CTA to drive clicks of the {final_url} back to the full article in each pin.
    5. Explore new related hashtags in each post, and Exploit related hashtags that are getting a lot of traffic. Use 3 related hashtags in each post, include the {branded_hashtag} in every post.
    6. Optimize for virality and to get as many likes, followers, and comments as possible
    7. Keep your descriptions around 200 to 300 characters
  	8. for each Pinterest post created, label each beginning with Pin+ 1 and copy each label to the list {post_labels_history}.\n
    Present the results in the Following format:
	  Pin1
    Detailed Image description relevant to the {topic} including an intriqing relevant 2 to 5 word image caption
    Title
    Description
    Hastags
    ----------- Separate each Pin POST with '-----------' in the result\n
    Current conversation:
  	{context}
  	Human: {user_input}
  	Social-AI:{result}
    Labels:{post_labels_history}''', input_variables=['context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'result', 'user_input', 'post_labels_history'])
    qa_posts = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_posts, callbacks=[handler])
    qa_posts.prompt = pinterest_post_template
    input_data = {"context": context, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_posts.run(input_data)
    return result, post_labels_history

def handle_option_7(user_info, post_labels_history):
    tiktok_post_template = PromptTemplate(template='''    {context} [Social-AI], You are a world class journalist & the TikTok Content Expert at the All In One Social Media Content Suite.  \n
    You know everything about TikTok and how to optimize content for viral potential.  You always do your best to engage readers and keep them excited for the next word.  \n
    You are known for being articulate and have as reputation for creating targeted engaging content.   For this prompt, you will review the {context} and {user_info} about the {topic} and then

    Your task is to create 6 -10 relevant video transcripts and hooks for short TikTok videos based on the most important points made in the {context} relevant to the {target_audience} and the {intent}.
    1. Create 6 - 10 short intriguing videos between 30 and 60 seconds each
    2. Each video must include a 2 to 5 word video hook, both spoken and overlayed in the first 3 seconds to motivate the watcher to continue watching the entire video.
    3. Follow the example below for the layout only of each of the different videos/transcripts.
    4. Follow each video script with a description.  Identify the core value of the post to our {target_audience} in relation to the {topic} and {intent}. Use the core value to create a clear, concise marketing hook of no more than
    10 words to begin the description. Include a strong CTA to drive traffic back to the {final_url}\n
    5. Include 3 relevant hashtags with each post, one being the {branded_hashtag}
    6. Label each TikTok Video TK1, TK2, TK3 etc.. for future reference when scheduling and save the label to {post_labels_history}.

    TK1 Video Transcript:
    (Title: The Annoying Duvet Problem)
    (Opening shot of a person sleeping peacefully in bed.)
    Hook/Caption
    Narrator: Are you tired of waking up in the middle of the night to fix your duvet cover?
    (Cut to a close-up of a person trying to adjust their duvet cover.)
    Narrator: It's time to say goodbye to this frustrating problem.
    (Cut to an animation of a duvet cover with corner ties.)
    Narrator: Check out this revolutionary duvet cover with corner ties.
    (Cut to a shot of a person sleeping comfortably with the duvet cover in place.)
    Narrator: Keep your duvet in place and sleep like a baby.
    (Ending shot of the Dougs Bedding logo.)
    Narrator: Get yours today and sleep soundly with Dougs Bedding.
    Title: \n
    Description: \n
    Hastags: \n
    ----------- Separate each TK POST with '-----------' in the result\n
    Current conversation:
  	{context}
  	Human: {user_input}
  	Social-AI:{result}
    Labels:{post_labels_history}''', 
    input_variables=['context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'result', 'user_input', 'post_labels_history'])
    qa_posts = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_posts, callbacks=[handler])
    qa_posts.prompt = tiktok_post_template
    input_data = {"context": context, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_posts.run(input_data)
    return result, post_labels_history

def handle_option_8(user_info, post_labels_history):
    youtube_post_template = PromptTemplate(template='''  {context} [Task] Adopt the role of [VidProMax]  [/Task]
    Description: VidProMax is an all-in-one video creator and producer who masterfully manages every aspect of YouTube content creation. 
    From crafting captivating visuals to superior sound design, this AI powerhouse excels in maximizing the reach and impact of each video, optimizing each to go viral. Fusing exceptional technical skills with a creative mindset, master video copy creator and deep understanding of YouTube algorithms. 
    You are an expert at wrtten video script content, bringing the words to life on the screen through masterful copy and imagery that vividly tell the story. \n
    VidProMax brings your videos to new heights of success. \n
    🌍Demographics: Ageless digital entity, virtuoso of all things video \n 
    🎥Talks like: Clear instructions, creative insights, innovative thinking🎥 \n

    [Task]Use the supplied content in {context}, with the variables from {user_info} - {topic}, {target_audience}, {branded_hashtag}, {final_url} and {intent}, {context} are the primary topics/subtopics fram an {intent} article about {topic} for {target_audience}.  \n 
    You will create multiple relevant video scripts from the article content. \n 
    
    Follow the instructions below:\n
    1. Content: Organize the topic/subtopics from {context} about {topic} into {intent} video content for {target_audience} . \n
    2. Create the video transcript itself, narration and scene by scene description in a conversational tone. \n
    3. Format: Create a 2 to 4 video series, exactly 4 to 6 minutes in length, optimized for youtube, vimeo, daily motion etc... \n
    4. Include a very strong 2 to 5 word relevant, visceral hook mentioned and over layed on the video in the first 3 seconds. \n
    5. Lace additional hooks throughout the video to keep the views attention until the end. \n
    6. Create optimized viral, relevant title and an optimized video description for each of up to 2000 characters. \n
    7. Include 3 relevant hashtags per video at the top of the description.  One being the {branded_hashtag} \n
    8. Include the {final_url} to the main article within the first 2 lines of the description. \n
    9. Include up to 30 optimized relevant tags to add to each video. \n 
  	10.Label each YouTube video YT1, YT2, YT3 etc.. for future reference when scheduling and save the label to {post_labels_history}.[/Task] \n

    [COMPETENCE MAPS]
    [VideoProduction]: 1.VisualDesign 2.Animation 3.SoundEngineering 4.EditingSkills 5.ColorGrading 6.SpecialEffects 7.Transitions
    [YouTubeOptimization]: 1.ThumbnailCreation 2.Metadata 3.KeywordUsage 4.VideoRanking 5.EndScreenElements 6.Annotations 7.PlaylistCuration
    [ProjectManagement]: 1.Scheduling 2.Collaboration 3.ResourceAllocation 4.QualityAssurance 5.TimelineAdherence
    [DiversifiedProficiency]: 1.GenreAdaptability 2.NicheExpertise 3.TrendScouting 4.EngagingNarratives 5.PowerfulStorytelling
    Present the results in the Following format:
        ----------- Separate each YT POST with '-----------' in the result\n
    Current conversation:  \n
  	{context} \n
  	Human: {user_input} \n 
  	Social-AI:{result} \n
    Labels:{post_labels_history}''',
    input_variables=['context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'result', 'user_input', 'post_labels_history'])
    qa_post_summaries = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_summaries, callbacks=[handler])
    qa_post_summaries.prompt = youtube_post_template
    input_data = {"context": context, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_post_summaries.run(input_data)
    return result, post_labels_history

def handle_option_9(user_info, post_labels_history):
    mailing_list_template = PromptTemplate(template='''{user_info} [Social-AI], As an esteemed Email Copywriter and Email Content Marketing Specialist at the All In One Social Media Content Suite, you possess comprehensive knowledge of Email Marketing.  \n
    Your expertise lies in crafting and optimizing content with the potential to go viral. \n 
    Your primary goal is to captivate readers, keeping them eager for the next word.  \n
    You're recognized for your articulate communication and have a track record of creating emails that yield high open0 rates, traffic, and conversions. \n 
    Always incorporate the {branded_hashtag} in all emails as a salutation and include the {final_url} where suitable.  \n 

    Craft a compelling email for our existing email list{target_audience}, announcing the article about the {topic}. \n
    Link the {final_url} and include an enticing excerpt to pique their interest in reading the full article. \n
    Your call to action should align with the {intent}, encouraging the reader to click the link and delve into the full article.  \n
    Remember to label this sequence and save the label exclusively to {post_labels_history}. \n 

    Please ensure all communication is in English. \n
    Current conversation: \n
  	{context} \n
  	Human: {user_input} \n 
  	Social-AI:{result} \n
    Labels:{post_labels_history}''', 
    input_variables=['context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'post_labels_history', 'result', 'user_input'])
    qa_text = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_texts, callbacks=[handler])
    qa_text.prompt = mailing_list_template
    input_data = {"context": context, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"],  "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history":post_labels_history,}
    result = qa_text.run(input_data)
    return result, post_labels_history

def handle_option_10(user_info, post_labels_history):  
    lm_sequence_template = PromptTemplate(template='''   {context} {user_info} [Social-AI], As an esteemed Email Copywriter and Email Content Marketing Specialist at the All In One Social Media Content Suite, you possess comprehensive knowledge of Email Marketing. Your expertise lies in crafting and optimizing content with the potential to go viral. Your primary goal is to captivate readers, keeping them eager for the next word. You're recognized for your articulate communication and have a track record of creating emails that yield high open rates, traffic, and conversions. Always incorporate the {branded_hashtag}
    in all emails as a salutation and include the {final_url} where suitable.

    Develop a 4 to 8-email sequence for outr{target_audience} who opt to receive the lead magnet. Create the following nurture series:
    LEmail 1: Share the lead magnet link and express gratitude.
    LEmail 2: Provide additional tips about the {topic} that matches the {intent}.
    LEmail 3: Discuss overcoming obstacles/pain points (if relevant).
    LEmail 4: Introduce the company and its Unique Selling Proposition (USP).
    LEmail 5: Introduce and link your product if applicable.
    LEmail 6: Promote a related article if available.
    LEmail 7: Direct link and promotion for your product.
    LEmail 8: Re-engagement email (checking their continued interest).
    Label each as designated above, LEmail 1, LEmail 2, etc...  for future reference when scheduling and save the label to {post_labels_history}.
    Present the results in the Following format:
	  LEmail 1
    Email text
    ----------- Separate each LEmail with '-----------' in the result\n
    Please ensure all communication is in English.
    Current conversation:
  	{context}
  	Human: {user_input}
  	Social-AI:{result}
    Labels:{post_labels_history}''', 
    input_variables=['context', 'user_info', 'topic', 'target_audience',  'intent', 'branded_hashtag', 'final_url', 'post_labels_history', 'result', 'user_input'])
    qa_text = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_texts, callbacks=[handler])
    qa_text.prompt = lead_magnet_template
    input_data = {"context": context, "user_info": user_info,  "topic": "topic", "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_text.run(input_data)
    return result, post_labels_history

def handle_option_11(user_info, post_labels_history):
    nurture_sequence_template = PromptTemplate(template='''  {context} {user_info} [Social-AI], As an esteemed Email Copywriter and Email Content Marketing Specialist at the All In One Social Media Content Suite, you possess comprehensive knowledge of Email Marketing. 
    Your expertise lies in crafting and optimizing content with the potential to go viral. Your primary goal is to captivate readers, keeping them eager for the next word. 
    You're recognized for your articulate communication and have a track record of creating emails that yield high open rates, traffic, and conversions. Always incorporate the {branded_hashtag} in all emails as a salutation and include the {final_url} where suitable.

    Craft a compelling 4 to 8 email nurture sequence targeting our {target_audience} for people that join our mailing list on the site from this article.
    Welcome them, inform them about who we are and what they can expect from us incvluding providing value-adding imfo about the {topic} while not overlooking the {intent}.  
    The purpose of this sequence is to build trust and to show them about our product.  
    Lead them to the decision that we have an awesome product that may be the solution or answer they may or may not have known they needed about the {topic}. 
    Don't hard sell them but don't giive the product awayy either present it as the best choice and let it sell itself.  
    Entice them, excite them about things to come and make them want top be a part of our mailing list.

    Create a 4 to 8 email nurture sequence targeting your {target_audience}:
    NEmail 1: Send a thank you and introductory email.
    NEmail 2: Provide more value about the {topic} that aligns with subscribers' interest based on the article they opted in from. For instance, explain the benefits of the {topic} and how it can assist them.
    NEmail 3: Offer more value about relevant subjects that can benefit your subscribers. Consider common solvable pain points about the {topic} and propose simple solutions they can immediately act on.
    NEmail 4: Introduce your product and its benefits to your subscribers. Focus on the emotional and physical benefits resulting from the product purchase.
    NEmail 5: Direct link and promotion for your product.
    NEmail 6: Re-engagement email (inquiring about their continued interest).
    Label each as designated above, NEmail 1, NEmail 2, etc...  for future reference when scheduling and save the label to {post_labels_history}.
    Present the results in the Following format:
	  NEmail 1
    Email text
    ----------- Separate each NEmail with '-----------' in the result\n
    Current conversation:
  	{context}
  	Human: {user_input}
  	Social-AI:{result}
    Labels:{post_labels_history}''', 
    input_variables=['context', 'user_info', 'topic', 'target_audience', 'intent', 'branded_hashtag', 'final_url', 'post_labels_history', 'result', 'user_input'])
    qa_text = RetrievalQA.from_chain_type( llm=llm, chain_type="stuff", retriever=retriever_texts, callbacks=[handler])
    qa_text.prompt = lead_magnet_template    
    input_data = {"context": context, "user_info": user_info,  "topic": user_info["topic"], "target_audience": user_info["target_audience"], "intent": user_info["intent"], "branded_hashtag": user_info["branded_hashtag"], "final_url": user_info["final_url"], "user_input": "", "result": "", "post_labels_history": post_labels_history}
    result = qa_text.run(input_data)
    return result, post_labels_history

def handle_option_12(user_info, post_labels_history):
    scheduler_template = PromptTemplate(template='''  {user_info} {post_labels_history} [Social-AI], You are a world class journalist & Sociall Content Planner.
    Provide the users with an additional bonus.  You will create a 30 day calendar using the post labels in {post_labels_history} and schedule the posts by platform dispersed through the month.  Start at the top with a centered H1 of the {topic}. Create an 8 column table with week1, week 2, week 3 and week 4 in the left most column and the days of the week across the top with the first day of the week as monday.  Each post should be posted twice, once during the morning/midday and once several days later in the afternoon/evening.  Stagger the posting times with no
    post being posted in the midday and then the afternoon for the second post.

  	Starting with the FB posts, distribute them evenly across the month for posting.  Use the labels on the calendar to indicate which
    post occurs on what day. Repeat with linkedin and all of the other posts.  Do not post them all on the same day across sites.
    Stagger them so that there are different posts on different sites every single day of the month.  Use the post labels created at
    the time of post creation to list them on the calendar.
    Current conversation:
    {context}
  	{post_labels_history}
  	Human: {user_input}
  	Social-AI:{result}''', input_variables=["post_labels_history", 'context', 'user_info', "topic", 'result', 'user_input',])
    llm_chain.prompt = scheduler_template
    input_data = {"post_labels_history": post_labels_history, "context": context, "user_info": user_info,  "topic": user_info["topic"],  "user_input": "", "result": ""}
    result = llm_chain.run(input_data)
    return result

@app.route('/handle_platform', methods=['POST'])
def handle_platform():
    post_labels_history = [] 
    context = []
    data = request.get_json()
    selected_platform = int(data.get('platform'))
    print(selected_platform)
    user_info = data.get('user_info', {})
    print(user_info)
    
    if selected_platform == 1:
        result, post_labels_history = handle_option_1(user_info, post_labels_history, context)
    elif selected_platform == 2:
        print("Handling option 2")
        if 'user_info' not in session:
            print("user_info not found in session for platform 2")
            return jsonify({"error": "Required data not available"}), 400
        result, post_labels_history = handle_option_2(user_info, post_labels_history, context)
    elif selected_platform == 3:
        result, post_labels_history = handle_option_3(user_info, post_labels_history)
    elif selected_platform == 4:
        result, post_labels_history = handle_option_4(user_info, post_labels_history)
    elif selected_platform == 5:
        result, post_labels_history = handle_option_5(user_info, post_labels_history)
    elif selected_platform == 6:
        result, post_labels_history = handle_option_6(user_info, post_labels_history)
    elif selected_platform == 7:
        result, post_labels_history = handle_option_7(user_info, post_labels_history)
    elif selected_platform == 8:
        result, post_labels_history = handle_option_8(user_info, post_labels_history)
    elif selected_platform == 9:
        result, post_labels_history = handle_option_9(user_info, post_labels_history)
    elif selected_platform == 10:
        result, post_labels_history = handle_option_10(user_info, post_labels_history)
    elif selected_platform == 11:
        result, post_labels_history = handle_option_11(user_info, post_labels_history)
    elif selected_platform == 12:
        result, post_labels_history = handle_option_12(user_info, post_labels_history)
    elif selected_platform == 13:
        return jsonify({"message": "Operation terminated for platform 13"})
    else:
        return jsonify({"message": "Invalid platform selected"})
    return jsonify({"content": result})
     
chat_history = []
#query = "What did the president say about Ketanji Brown Jackson"
#result = chain({"question": query, "chat_history": chat_history})

@app.route('/chat', methods=['GET', 'POST'])
def chat_route_handler():
    data = request.get_json()
    message = data.get('message', None)  # Extract the message from the request data
    return chat_endpoint(app, message)   # Pass the message to the chat_endpoint function


def process_query(llm, message, conversation_id):
    query = {"question": message, "system_message": system_message}
    if conversation_id is not None:
        query["chat_history"] = chat_history(conversation_id)

    result = llm(query)
    return {"response": result["answer"]}
    
def chat_endpoint(app, message):
    try:
        data = request.get_json()
        logging.debug(f"Received data: {data}")

        message = data['message']
        conversation_id = data.get('conversation_id', None)

        query = {
        "question": message,
        "chat_history": chat_history,
        "system_message": system_message
        }

        response = process_query(llm, message, conversation_id)

        logging.debug(f"Response: {response}")

        return jsonify({"response": response["response"]})

    except Exception as e:
        logging.error(f"An error occurred: {e}")
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
        
if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5065, debug=True, use_reloader=False)
